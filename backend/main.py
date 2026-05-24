
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import httpx
import os
import json
import re
import tempfile
import shutil
import pdfplumber
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="Recruiter Email Generator API")

# ── CORS ──────────────────────────────────────────────────────────────────────
FRONTEND_URL = os.getenv("FRONTEND_URL", "")
ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:5173",
    "http://127.0.0.1:3000",
    "http://127.0.0.1:5173",
]
if FRONTEND_URL:
    ALLOWED_ORIGINS.append(FRONTEND_URL)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GROQ_API_KEY    = os.getenv("GROQ_API_KEY", "").strip()
GROQ_SM_API_KEY = os.getenv("GROQ_SM_API_KEY", "").strip()
GROQ_API_URL    = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL      = "llama-3.3-70b-versatile"

def get_sm_key() -> str:
    key = GROQ_SM_API_KEY or GROQ_API_KEY
    if not key:
        raise HTTPException(status_code=500, detail="No Groq API key configured.")
    return key

print(f"[STARTUP] GROQ_API_KEY    : {'SET' if GROQ_API_KEY else 'MISSING'}")
print(f"[STARTUP] GROQ_SM_API_KEY : {'SET' if GROQ_SM_API_KEY else 'MISSING (will use fallback)'}")


# ═══════════════════════════════════════════════════════════════════════════════
# MODELS
# ═══════════════════════════════════════════════════════════════════════════════

class GenerateSizzlingRequest(BaseModel):
    cv_text: str
    jd_text: str

class SizzlingResponse(BaseModel):
    candidate_name: str
    top_sentences: list[str]
    skills: list[str]


# ═══════════════════════════════════════════════════════════════════════════════
# SKILL FILTER
# ═══════════════════════════════════════════════════════════════════════════════

def skill_exists_in_cv(skill: str, cv_text: str) -> bool:
    STOP = {"and", "&", "/", "the", "with", "for", "in", "of", "a", "an"}
    cv_lower = cv_text.lower()
    words = [w.strip(".,()[]") for w in skill.lower().split()
             if w.strip(".,()[]") not in STOP]
    return all(word in cv_lower for word in words if word)

def filter_skills_to_cv(skills: list[str], cv_text: str) -> list[str]:
    return [s for s in skills if skill_exists_in_cv(s, cv_text)]


# ═══════════════════════════════════════════════════════════════════════════════
# GROQ CALLER
# ═══════════════════════════════════════════════════════════════════════════════

async def call_groq(prompt: str, api_key: str, label: str = "groq") -> str:
    if not api_key:
        raise HTTPException(status_code=500, detail="API key not configured.")
    print(f"[{label}] → Groq  model={GROQ_MODEL}  prompt={len(prompt)} chars")
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": GROQ_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens": 4096,
    }
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(GROQ_API_URL, headers=headers, json=payload)
        print(f"[{label}] ← Groq  status={resp.status_code}")
        if resp.status_code == 401:
            raise HTTPException(status_code=401, detail="Invalid Groq API key.")
        if resp.status_code == 429:
            w = resp.headers.get("retry-after", "20-30")
            raise HTTPException(status_code=429, detail=f"Rate limit. Wait {w}s.")
        if resp.status_code == 413:
            raise HTTPException(status_code=413, detail="Input too large for Groq.")
        if resp.status_code != 200:
            raise HTTPException(status_code=resp.status_code,
                detail=f"Groq error ({resp.status_code}): {resp.text[:300]}")
        content = resp.json()["choices"][0]["message"]["content"].strip()
        print(f"[{label}] ← {len(content)} chars returned")
        if not content:
            raise HTTPException(status_code=500, detail="Groq returned empty response.")
        return content
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Groq timed out.")
    except httpx.ConnectError:
        raise HTTPException(status_code=503, detail="Cannot reach Groq API.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unexpected: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# PII STRIPPER
# ═══════════════════════════════════════════════════════════════════════════════

def strip_cv_pii(text: str) -> tuple[str, dict]:
    found: dict[str, list[str]] = {
        "emails": [], "phones": [], "ids": [],
        "dobs": [], "addresses": [], "urls": [],
    }

    def sub_and_track(pattern, replacement, src, key, flags=0):
        matches = re.findall(pattern, src, flags)
        for m in matches:
            val = (m if isinstance(m, str) else m[0]).strip()
            if val and val not in found[key]:
                found[key].append(val)
        return re.sub(pattern, replacement, src, flags)

    text = sub_and_track(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
        "[EMAIL REDACTED]", text, "emails"
    )
    text = sub_and_track(
        r"\+?\d[\d\s\-().]{7,}\d",
        "[PHONE REDACTED]", text, "phones"
    )
    text = sub_and_track(
        r"\b\d{6,8}[-]\d{4}\b",
        "[ID REDACTED]", text, "ids"
    )
    text = sub_and_track(
        r"\b\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}\b",
        "[DOB REDACTED]", text, "dobs", flags=re.IGNORECASE
    )
    text = sub_and_track(
        r"(?:Passport|National ID|NIN|SSN|BSN|PAN|Aadhar|ID No\.?)\s*[:\-]?\s*[A-Z0-9\-]{5,20}",
        "[ID REDACTED]", text, "ids", flags=re.IGNORECASE
    )
    text = sub_and_track(
        r"\d+[a-zA-Z]?\s+[A-Z][a-zA-Z\s]{3,}(?:Street|St|Road|Rd|Avenue|Ave|Lane|Ln|Drive|Dr|Close|Crescent|Way|Place|Pl)",
        "[ADDRESS REDACTED]", text, "addresses", flags=re.IGNORECASE
    )
    text = sub_and_track(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/[^\s,;]+",
        "[LINKEDIN REDACTED]", text, "urls", flags=re.IGNORECASE
    )
    text = sub_and_track(
        r"https?://[^\s,;<>\"']{4,}",
        "[URL REDACTED]", text, "urls", flags=re.IGNORECASE
    )
    text = sub_and_track(
        r"www\.[^\s,;<>\"']{4,}",
        "[URL REDACTED]", text, "urls", flags=re.IGNORECASE
    )

    total = sum(len(v) for v in found.values())
    print(f"[PII] Stripped {total} items — CV safe for AI ({len(text)} chars)")
    return text, found


# ═══════════════════════════════════════════════════════════════════════════════
# JSON PARSERS
# ═══════════════════════════════════════════════════════════════════════════════

def _strip_fences(raw: str) -> str:
    raw = re.sub(r"^```(?:json)?\s*", "", raw.strip(), flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    return raw.strip()

def parse_json_response(raw: str) -> dict:
    raw = _strip_fences(raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    s, e = raw.find("{"), raw.rfind("}") + 1
    if s != -1 and e > s:
        try:
            return json.loads(raw[s:e])
        except json.JSONDecodeError:
            pass
    raise HTTPException(status_code=500, detail="AI returned unexpected format.")

def parse_json_array(raw: str, label: str = "sm") -> list:
    original = raw
    raw = _strip_fences(raw)
    try:
        result = json.loads(raw)
        if isinstance(result, list):
            print(f"[{label}] parsed directly: {len(result)} items")
            return result
    except json.JSONDecodeError:
        pass
    s = raw.find("[")
    if s != -1:
        e = raw.rfind("]")
        if e > s:
            try:
                result = json.loads(raw[s:e+1])
                if isinstance(result, list):
                    print(f"[{label}] parsed via bracket search: {len(result)} items")
                    return result
            except json.JSONDecodeError:
                pass
        arr = raw[s:]
        items, depth, obj_start = [], 0, None
        for i, ch in enumerate(arr):
            if ch == "{":
                if depth == 0:
                    obj_start = i
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0 and obj_start is not None:
                    fragment = arr[obj_start:i+1]
                    fragment = re.sub(r",\s*([}\]])", r"\1", fragment)
                    try:
                        obj = json.loads(fragment)
                        if isinstance(obj, dict) and "requirement" in obj:
                            items.append(obj)
                    except json.JSONDecodeError:
                        pass
                    obj_start = None
        if items:
            print(f"[{label}] repaired {len(items)} items from truncated array")
            return items
    print(f"[{label}] PARSE FAILED. First 500 chars:\n{original[:500]}")
    raise HTTPException(status_code=500,
        detail=f"AI format error. Starts with: {original[:120]!r}")


# ═══════════════════════════════════════════════════════════════════════════════
# SIZZLING
# ═══════════════════════════════════════════════════════════════════════════════

def build_sizzling_prompt(cv: str, jd: str) -> str:
    return f"""Act as a senior recruiter. You will receive a candidate CV and a Job Description (JD).
Perform ALL tasks and return a single valid JSON object. No markdown, no extra text, JSON only.

TASK 1 — Extract from CV:
- candidate_name: The candidate's full name
- designation: Their current or most recent job title/role (e.g., "Full Stack Developer", "Test Manager", "Data Engineer")
- years_experience: Total years of professional experience (e.g., "4+", "10+", "15+")

TASK 2 — top_sentences: Generate exactly 3 sentences matching CV to JD.

SENTENCE 1 FORMAT (MANDATORY):
"[Candidate Name] is a [Designation] with [X]+ years of experience in [key skills/domains from CV that match JD], with hands-on expertise in [specific tools/technologies from CV]."

SENTENCE 2 FORMAT:
Start with "Proficient in" or "Skilled in" — describe specific technical capabilities, tools, and methodologies from the CV that match JD requirements.

SENTENCE 3 FORMAT:
Start with "Demonstrated" or "Adept at" — describe accomplishments, domain experience, or measurable achievements from the CV relevant to JD.

Rules for all sentences:
- Each sentence MUST be 2-3 lines long with specific details.
- Use ONLY information that EXISTS in the CV. Do not invent anything.
- Focus on skills/experience matching the JD requirements.

TASK 3 — skills: Extract exactly 15 skills from the CV most relevant to the JD.
Rules:
- ALL skills must exist word-for-word in the CV. Do not invent any skill.
- Prioritize skills mentioned in the JD that also appear in the CV.
- Short phrases only (1-4 words each).

Return ONLY this JSON (no markdown, no extra text):
{{
  "candidate_name": "Full Name",
  "designation": "Job Title",
  "years_experience": "X+",
  "top_sentences": [
    "Sentence 1 starting with candidate name.",
    "Sentence 2 starting with Proficient/Skilled.",
    "Sentence 3 starting with Demonstrated/Adept."
  ],
  "skills": ["Skill 1","Skill 2","Skill 3","Skill 4","Skill 5","Skill 6","Skill 7","Skill 8","Skill 9","Skill 10","Skill 11","Skill 12","Skill 13","Skill 14","Skill 15"]
}}

CV:
{cv}

JD:
{jd}

JSON:"""


@app.post("/api/generate-sizzling", response_model=SizzlingResponse)
async def generate_sizzling(request: GenerateSizzlingRequest):
    cv = request.cv_text.strip()
    jd = request.jd_text.strip()
    if not cv:
        raise HTTPException(status_code=400, detail="CV text is required.")
    if not jd:
        raise HTTPException(status_code=400, detail="Job description is required.")
    if len(cv) > 40_000:
        raise HTTPException(status_code=413, detail="CV too long (>40,000 chars).")
    if len(jd) > 10_000:
        raise HTTPException(status_code=413, detail="JD too long (>10,000 chars).")

    raw = await call_groq(build_sizzling_prompt(cv, jd), GROQ_API_KEY, label="sizzling")
    parsed = parse_json_response(raw)
    raw_skills: list[str] = (parsed.get("skills", []) or [])[:15]
    verified = filter_skills_to_cv(raw_skills, cv)
    dropped = [s for s in raw_skills if s not in verified]
    if dropped:
        print(f"[sizzling] Dropped hallucinated skills: {dropped}")
    return SizzlingResponse(
        candidate_name=str(parsed.get("candidate_name", "")).strip().strip('"'),
        top_sentences=(parsed.get("top_sentences", []) or [])[:3],
        skills=verified,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# FILE TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

_OLE_GARBAGE_PATTERNS = re.compile(
    r"\[Content_Types\]|_rels/|theme/theme|\.xmlPK|"
    r"<\?xml|<a:clrMap|Normal\.dotm|CompObj|Root Entry|"
    r"SummaryInformation|DocumentSummary|WordDocument stream|"
    r"Microsoft Office Word|KSOProduct|ContentTypeId|"
    r"eyJoZGlk|0x0101|Word\.Document\.",
    re.IGNORECASE
)

_JUNK_LINE_PATTERNS = re.compile(
    r"^\[Content_Types\]|^_rels/|^theme/|\.xmlPK$|"
    r"^<\?xml|^<a:|Normal\.dotm|^Root Entry$|^CompObj$|"
    r"^WordDocument$|^1Table$|^0Table$|^Data$|"
    r"^SummaryInformation|^DocumentSummary|"
    r"Microsoft Office Word|^KSO|^ContentType|"
    r"^eyJ|^0x01|Word\.Document\.|^\d{4}-\d{2}\.",
    re.IGNORECASE
)

def _is_garbage_line(line: str) -> bool:
    line = line.strip()
    if not line:
        return True
    if _JUNK_LINE_PATTERNS.search(line):
        return True
    if re.match(r'^[A-Za-z0-9+/=]{40,}$', line):
        return True
    if re.match(r'^[0-9A-Fa-f]{20,}$', line):
        return True
    if line.startswith("<") and ("xmlns" in line or "://" in line):
        return True
    return False

def _clean_doc_text(text: str) -> str:
    lines = text.splitlines()
    clean = []
    for line in lines:
        if _OLE_GARBAGE_PATTERNS.search(line):
            print(f"[DOC-CLEAN] Stopped at OLE sentinel: {line[:60]!r}")
            break
        if not _is_garbage_line(line):
            clean.append(line)
    return "\n".join(clean).strip()

def _read_docx(path: str) -> str:
    from docx import Document as D
    doc = D(path)
    seen, parts = set(), []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and t not in seen:
            seen.add(t)
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    parts.append(t)
    return "\n".join(parts)

def _read_doc_binary(path: str) -> str:
    try:
        text = _read_docx(path)
        if text.strip():
            print(f"[DOC] S1 (python-docx): {len(text)} chars")
            return _clean_doc_text(text)
    except Exception as e:
        print(f"[DOC] S1 failed: {e}")

    try:
        with open(path, "rb") as f:
            data = f.read()
        chunks, i = [], 0
        while i < len(data) - 1:
            word = int.from_bytes(data[i:i+2], "little")
            if 32 <= word <= 126 or word in (9, 10, 13):
                run = []
                j = i
                while j < len(data) - 1:
                    w = int.from_bytes(data[j:j+2], "little")
                    if 32 <= w <= 126 or w in (9, 10, 13, 160):
                        run.append(chr(w))
                        j += 2
                    else:
                        break
                if len(run) >= 5:
                    chunk = "".join(run).strip()
                    alpha = sum(1 for c in chunk if c.isalpha())
                    if chunk and alpha / max(len(chunk), 1) > 0.35:
                        chunks.append(chunk)
                i = j if j > i else i + 2
            else:
                i += 2
        text = "\n".join(chunks)
        cleaned = _clean_doc_text(text)
        if cleaned.strip():
            print(f"[DOC] S2 (UTF-16LE scan): {len(cleaned)} chars")
            return cleaned
    except Exception as e:
        print(f"[DOC] S2 failed: {e}")

    try:
        with open(path, "rb") as f:
            data = f.read()
        runs = re.findall(rb"[ -~]{6,}", data)
        lines = []
        for r in runs:
            s = r.decode("ascii", errors="ignore").strip()
            words = re.findall(r"[a-zA-Z]{3,}", s)
            if len(words) >= 2 and len(s) > 8:
                lines.append(s)
        text = "\n".join(lines)
        cleaned = _clean_doc_text(text)
        if cleaned.strip():
            print(f"[DOC] S3 (ASCII scan): {len(cleaned)} chars")
            return cleaned
    except Exception as e:
        print(f"[DOC] S3 failed: {e}")

    raise HTTPException(
        status_code=500,
        detail=(
            "Cannot read this .doc file. Please open it in Word, "
            "save as .docx (File → Save As → .docx), then re-upload."
        )
    )

def extract_file_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    print(f"[EXTRACT] {ext}  path={path}")

    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()

    if ext == ".docx":
        try:
            return _read_docx(path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to read DOCX: {e}")

    if ext == ".doc":
        return _read_doc_binary(path)

    if ext == ".pdf":
        try:
            pages = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            return "\n".join(pages).strip()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to read PDF: {e}")

    raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")


# ═══════════════════════════════════════════════════════════════════════════════
# CV PREVIEW ENDPOINT  ← NEW
# Extracts text from the uploaded CV, strips PII, returns safe text + PII report
# so the frontend can show the redacted CV preview BEFORE generation starts.
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/preview-cv")
async def preview_cv(cv: UploadFile = File(...)):
    name = (cv.filename or "").lower()
    if not any(name.endswith(e) for e in (".docx", ".doc", ".txt", ".pdf")):
        raise HTTPException(status_code=400,
            detail="CV must be .docx, .doc, .txt, or .pdf.")

    tmp_dir = tempfile.mkdtemp()
    try:
        ext      = Path(cv.filename or "cv.docx").suffix.lower()
        cv_path  = os.path.join(tmp_dir, f"cv{ext}")
        cv_bytes = await cv.read()
        with open(cv_path, "wb") as f:
            f.write(cv_bytes)

        # Extract raw text
        raw_text = extract_file_text(cv_path)
        if not raw_text.strip():
            raise HTTPException(status_code=400,
                detail="Could not extract text from CV. Try .docx or .txt format.")

        # Strip PII
        safe_text, pii_found = strip_cv_pii(raw_text)

        # Normalise whitespace (saves tokens + cleans up PDF artefacts)
        safe_text_clean = "\n".join(
            line.strip() for line in safe_text.splitlines() if line.strip()
        )

        return {
            "safe_text":     safe_text_clean,
            "raw_chars":     len(raw_text),
            "safe_chars":    len(safe_text_clean),
            "pii": {
                "emails":    pii_found.get("emails",    []),
                "phones":    pii_found.get("phones",    []),
                "ids":       pii_found.get("ids",       []),
                "dobs":      pii_found.get("dobs",      []),
                "addresses": pii_found.get("addresses", []),
                "urls":      pii_found.get("urls",      []),
            },
        }
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════════════════════
# SKILL MATRIX ROW EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

_HEADER_SKIP = {
    "requirements", "requirement",
    "explain how the candidate meets the requirement",
    "how the candidate meets the requirement",
    "skill", "skills", "criteria", "competency", "competencies",
    "communication",
}

_ROW_JUNK = re.compile(
    r"\[Content_Types\]|_rels/|\.xmlPK|<\?xml|<a:|Normal\.dotm|"
    r"eyJ[A-Za-z0-9]|^0x01|Word\.Document|KSOProduct|ContentType",
    re.IGNORECASE
)

def _is_valid_requirement(text: str) -> bool:
    t = text.strip()
    if not t or len(t) < 5:
        return False
    if t.lower() in _HEADER_SKIP:
        return False
    if _ROW_JUNK.search(t):
        return False
    if re.match(r'^[A-Za-z0-9+/=]{30,}$', t):
        return False
    if re.match(r'^[0-9A-Fa-f]{16,}$', t):
        return False
    if not re.search(r'[a-zA-Z]{3,}', t):
        return False
    return True

def extract_skill_matrix_rows(path: str) -> list[str]:
    ext = Path(path).suffix.lower()

    if ext == ".docx":
        try:
            from docx import Document as D
            doc = D(path)
            rows_found = []
            for table in doc.tables:
                if not table.rows:
                    continue
                first_cell = table.rows[0].cells[0].text.strip().lower()
                skip_header = any(
                    kw in first_cell
                    for kw in ("requirement", "skill", "criteria", "competenc", "explain")
                )
                for i, row in enumerate(table.rows):
                    if i == 0 and skip_header:
                        continue
                    cell = row.cells[0].text.strip()
                    cell = re.sub(r"\*\*(.+?)\*\*", r"\1", cell)
                    if _is_valid_requirement(cell):
                        rows_found.append(cell)
                if rows_found:
                    break
            if rows_found:
                print(f"[SM-ROWS] {len(rows_found)} rows from DOCX table")
                return rows_found
        except Exception as e:
            print(f"[SM-ROWS] DOCX table read failed: {e}")

    plain = extract_file_text(path)
    rows = _parse_plain_text_rows(plain)
    print(f"[SM-ROWS] {len(rows)} rows from plain text ({ext})")
    return rows

def _parse_plain_text_rows(text: str) -> list[str]:
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if _OLE_GARBAGE_PATTERNS.search(line):
            print(f"[PARSE-ROWS] Stopped at OLE sentinel: {line[:60]!r}")
            break
        if not line or not _is_valid_requirement(line):
            continue
        if line.startswith("|"):
            if "---" in line:
                continue
            parts = [p.strip() for p in line.strip("|").split("|")]
            cell = re.sub(r"\*\*(.+?)\*\*", r"\1", parts[0]).strip() if parts else ""
            if _is_valid_requirement(cell):
                rows.append(cell)
            continue
        m = re.match(r"^[-•*]\s+(.+)$", line)
        if m:
            candidate = m.group(1).strip()
            if _is_valid_requirement(candidate):
                rows.append(candidate)
            continue
        m = re.match(r"^\d+[.)]\s+(.+)$", line)
        if m:
            candidate = m.group(1).strip()
            if _is_valid_requirement(candidate):
                rows.append(candidate)
            continue
        if len(line) > 10 and _is_valid_requirement(line):
            rows.append(line)
    return rows


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION HEADER DETECTOR
# ═══════════════════════════════════════════════════════════════════════════════

SECTION_KEYWORDS = {
    "must have", "nice to have", "individual quality", "should have",
    "required", "preferred", "mandatory", "desirable", "good to have",
}

def is_section_header(text: str) -> bool:
    t = text.strip().lower().rstrip(":")
    return t in SECTION_KEYWORDS or (len(text.strip()) < 50 and text.strip().endswith(":"))


# ═══════════════════════════════════════════════════════════════════════════════
# SKILL MATRIX PROMPT
# ═══════════════════════════════════════════════════════════════════════════════

def build_skill_matrix_prompt(requirements: list[str], cv_text: str) -> str:
    cv_trimmed = cv_text[:7000]
    if len(cv_text) > 7000:
        print(f"[SM-PROMPT] CV trimmed {len(cv_text)} → 7000 chars")
    req_lines = "\n".join(f"{i+1}. {r}" for i, r in enumerate(requirements))

    return f"""You are a senior recruiter filling in a Skill Matrix. For each requirement, write how the candidate meets it using ONLY information from the CV provided.

CRITICAL RULES — follow every one exactly:
1. Return ONLY a valid JSON array. No markdown fences, no text before or after the array.
2. Each array item: {{"requirement": "<copy requirement text exactly>", "answer": "<your answer>"}}
3. For section header rows (e.g. "Must have:", "Nice to have:", "Individual quality:") set answer to "".
4. NEVER mention the candidate by name anywhere in any answer.
5. Start every non-empty answer with one of these strong words: Experienced, Proficient, Skilled, Adept, Demonstrated, Proven, Accomplished, Competent, Versed, Specialised.
6. Every keyword, tool, technology, and achievement in the answer MUST come from the CV — do not invent or infer anything.
7. If the CV does not support a requirement, write exactly: "Limited evidence in CV."
8. Keep each answer to 2-3 sentences. Professional, confident, third-person tone.
9. The "requirement" value must be copied EXACTLY from the numbered list below (without the number prefix).

REQUIREMENTS (fill answers for each):
{req_lines}

CANDIDATE CV (your only evidence source):
{cv_trimmed}

Output the JSON array now (starting with [ ):"""


# ═══════════════════════════════════════════════════════════════════════════════
# SKILL MATRIX DOCX BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_skill_matrix_docx(rows: list[dict], output_path: str):
    from docx import Document as D
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(old)
        b = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "595959")
            b.append(el)
        tcPr.append(b)

    def set_width(cell, twips: int):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcW")):
            tcPr.remove(old)
        w = OxmlElement("w:tcW")
        w.set(qn("w:w"), str(twips))
        w.set(qn("w:type"), "dxa")
        tcPr.append(w)

    def write_cell(cell, text: str, bold=False, size=10,
                   align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text or "")
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)

    doc = D()
    sec = doc.sections[0]
    sec.page_width    = Cm(29.7)
    sec.page_height   = Cm(21.0)
    sec.left_margin   = Cm(1.5)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    COL1, COL2 = 4252, 10886

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    write_cell(hdr[0], "Requirements",
               bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(hdr[1], "Explain how the candidate meets the requirement",
               bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in hdr:
        set_bg(c, "DEEAF6")
        set_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_width(hdr[0], COL1)
    set_width(hdr[1], COL2)

    for rd in rows:
        is_hdr = rd.get("is_header", False)
        nr = table.add_row().cells
        write_cell(nr[0], rd.get("requirement", ""), bold=is_hdr, size=10)
        write_cell(nr[1], rd.get("answer", ""),      bold=False,  size=10)
        for c, bg in [(nr[0], "F2F2F2" if is_hdr else "FFFFFF"),
                      (nr[1], "FFFFFF")]:
            set_bg(c, bg)
            set_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_width(nr[0], COL1)
        set_width(nr[1], COL2)

    doc.save(output_path)
    print(f"[SM-DOCX] Saved → {output_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# SKILL MATRIX ENDPOINT
# ═══════════════════════════════════════════════════════════════════════════════

@app.post("/api/generate-skill-matrix")
async def generate_skill_matrix(
    skill_matrix: UploadFile = File(...),
    cv: UploadFile = File(...),
):
    for upload, label in [(skill_matrix, "Skill Matrix"), (cv, "CV")]:
        name = (upload.filename or "").lower()
        if not any(name.endswith(e) for e in (".docx", ".doc", ".txt", ".pdf")):
            raise HTTPException(status_code=400,
                detail=f"{label} must be .docx, .doc, .txt, or .pdf.")

    tmp_dir = tempfile.mkdtemp()
    try:
        sm_ext  = Path(skill_matrix.filename or "sm.docx").suffix.lower()
        cv_ext  = Path(cv.filename or "cv.docx").suffix.lower()
        sm_path = os.path.join(tmp_dir, f"skill_matrix{sm_ext}")
        cv_path = os.path.join(tmp_dir, f"cv{cv_ext}")

        sm_bytes = await skill_matrix.read()
        cv_bytes = await cv.read()
        with open(sm_path, "wb") as f: f.write(sm_bytes)
        with open(cv_path, "wb") as f: f.write(cv_bytes)
        print(f"[SM] SM={len(sm_bytes)}B  CV={len(cv_bytes)}B")

        sm_rows = extract_skill_matrix_rows(sm_path)
        print(f"[SM] {len(sm_rows)} requirement rows: {sm_rows}")

        if not sm_rows:
            raise HTTPException(status_code=400,
                detail=(
                    "No requirement rows found in the Skill Matrix. "
                    "Ensure it has a 2-column table with requirements in the left column. "
                    "If it's a .doc file, resave as .docx in Word first."
                ))

        cv_text = extract_file_text(cv_path)
        print(f"[SM] CV text: {len(cv_text)} chars")
        if not cv_text.strip():
            raise HTTPException(status_code=400,
                detail="Could not extract text from CV. Try .docx or .txt format.")

        candidate_name = "Candidate"
        for line in cv_text.splitlines():
            line = line.strip()
            if re.match(r"^[A-Z][a-z]+(?: [A-Z][a-z]+){1,3}$", line):
                candidate_name = line
                break
        if candidate_name == "Candidate":
            m = re.search(r"^([A-Z][a-z]+(?: [A-Z][a-z]+)+)", cv_text, re.MULTILINE)
            if m:
                candidate_name = m.group(1).strip()
        print(f"[SM] Candidate (filename): {candidate_name}")

        cv_text_safe, pii_report = strip_cv_pii(cv_text)
        cv_safe_clean = "\n".join(
            line.strip() for line in cv_text_safe.splitlines() if line.strip()
        )

        prompt  = build_skill_matrix_prompt(sm_rows, cv_safe_clean)
        raw     = await call_groq(prompt, get_sm_key(), label="skill-matrix")
        ai_rows = parse_json_array(raw, label="skill-matrix")

        if not ai_rows:
            raise HTTPException(status_code=500,
                detail="AI returned an empty skill matrix. Please try again.")

        print(f"[SM] AI filled {len(ai_rows)} rows")

        final_rows = []
        for item in ai_rows:
            if not isinstance(item, dict):
                continue
            req = str(item.get("requirement", "")).strip()
            ans = str(item.get("answer", "")).strip()
            if candidate_name and candidate_name != "Candidate":
                ans = ans.replace(candidate_name, "The candidate")
            final_rows.append({
                "requirement": req,
                "answer":      ans,
                "is_header":   is_section_header(req),
            })

        safe  = re.sub(r"[^\w\s-]", "", candidate_name).strip().replace(" ", "_")
        fname = f"Skill_Matrix_{safe}.docx"
        out   = os.path.join(tmp_dir, fname)
        build_skill_matrix_docx(final_rows, out)

        pdir  = tempfile.mkdtemp(prefix="sm_out_")
        ppath = os.path.join(pdir, fname)
        shutil.copy2(out, ppath)

        pii_summary = {
            "emails":    pii_report.get("emails", []),
            "phones":    pii_report.get("phones", []),
            "ids":       pii_report.get("ids", []),
            "dobs":      pii_report.get("dobs", []),
            "addresses": pii_report.get("addresses", []),
            "urls":      pii_report.get("urls", []),
            "cv_safe_preview": cv_safe_clean[:800],
            "cv_safe_chars":   len(cv_safe_clean),
            "cv_raw_chars":    len(cv_text),
        }
        import base64 as _b64
        pii_header = _b64.b64encode(
            json.dumps(pii_summary, ensure_ascii=False).encode()
        ).decode()

        return FileResponse(
            path=ppath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=fname,
            headers={"X-PII-Report": pii_header},
        )

    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════════════════════
# HEALTH CHECK
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "groq_configured":    bool(GROQ_API_KEY),
        "groq_sm_configured": bool(GROQ_SM_API_KEY),
        "groq_sm_source":     "GROQ_SM_API_KEY" if GROQ_SM_API_KEY else "GROQ_API_KEY (fallback)",
        "model":              GROQ_MODEL,
    }